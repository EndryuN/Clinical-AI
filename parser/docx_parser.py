"""
Document parser for MDT outcome proformas (.docx).

The Word document contains exactly one table per patient (50 patients total).
Each table has a consistent 8-row x 3-col structure:
  Row 0: headers (Patient Details / Cancer Target Dates)
  Row 1: patient demographics (hospital number, NHS number, name, gender, DOB)
  Row 2: staging/diagnosis header
  Row 3: diagnosis + staging detail
  Row 4: clinical details header
  Row 5: clinical details free text
  Row 6: MDT outcome header
  Row 7: MDT outcome free text

Columns 1 and 2 are often duplicates of column 0; we only use column 0 (and
column 2 for cancer target dates where appropriate).
"""

import re
import sys
import os
from typing import Optional

# Allow imports from project root regardless of how this module is invoked
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from docx import Document  # type: ignore
from models import PatientBlock, CellRef


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

_NHS_RE = re.compile(r'NHS Number:\s*([\d\s()\w]+?)(?:\n|$)', re.IGNORECASE)
_HOSPITAL_RE = re.compile(r'Hospital Number:\s*(\S+)', re.IGNORECASE)
_GENDER_RE = re.compile(r'\b(Male|Female)\b', re.IGNORECASE)
# Name may appear as:
#   "AIDEN O'CONNOR(b)"          – all-caps with trailing annotation
#   "Erin Hall"                   – mixed-case after hospital/NHS lines
#   "Name: Noah Robin"            – explicit "Name:" prefix
_NAME_RE = re.compile(
    r'(?:Name:\s*)?([A-Z][A-Za-z\'\-]+(?:\s+[A-Za-z\'\-]+)+)',
)


def _clean(text: str) -> str:
    """Remove annotation markers like (a), (b), (c)... from cell text."""
    return re.sub(r'\([a-z]\)', '', text).strip()


def _extract_name(details_text: str) -> str:
    """
    Return the patient name found in the details cell, or empty string.

    Strategies tried in order:
    1. Explicit 'Name: <value>' prefix
    2. A line that is all-uppercase (old-style records)
    3. The third non-empty line (after Hospital Number and NHS Number lines)
    """
    # Strategy 1: explicit prefix
    name_prefix = re.search(r'Name:\s*([^\n]+)', details_text)
    if name_prefix:
        return _clean(name_prefix.group(1))

    lines = [ln.strip() for ln in details_text.splitlines() if ln.strip()]
    # Strategy 2: all-caps line (could have annotation suffix like "(b)")
    # Skip lines that look like medical data, not names
    _NOT_NAME = {'DAY', 'TARGET', 'BREACH', 'DATE', 'PATHWAY', 'PLEASE', 'TREATMENT',
                 'DECISION', 'STAGING', 'DIAGNOSIS', 'CLINICAL', 'MDT', 'OUTCOME',
                 'NOTE', 'NUMBER', 'HOSPITAL', 'NHS'}
    for line in lines:
        cleaned = _clean(line)
        if cleaned and cleaned.replace(' ', '').replace("'", '').replace('-', '').isupper():
            words = cleaned.upper().split()
            # If any word is a known non-name keyword, skip this line
            if not any(w in _NOT_NAME for w in words):
                return cleaned

    # Strategy 3: third non-empty line
    # lines[0] ~ "Hospital Number: ..."
    # lines[1] ~ "NHS Number: ..."
    # lines[2] ~ name (mixed-case or capitalised)
    if len(lines) >= 3:
        candidate = _clean(lines[2])
        # Sanity check: must look like a name (contains a space, no digits)
        if ' ' in candidate and not re.search(r'\d', candidate):
            return candidate

    return ""


def _initials(name: str) -> str:
    """Convert 'AIDEN O CONNOR' or 'Erin Hall' to 'AO' / 'EH'."""
    parts = re.split(r"[\s'\-]+", name.strip())
    return "".join(p[0].upper() for p in parts if p)


def _extract_nhs(details_text: str) -> str:
    """Return NHS number (digits only) or empty string."""
    m = _NHS_RE.search(details_text)
    if not m:
        return ""
    raw = m.group(1)
    # Strip annotation markers and whitespace
    digits = re.sub(r'[^\d]', '', raw)
    return digits


def _extract_gender(details_text: str) -> str:
    """Return 'Male', 'Female', or '' if not found."""
    m = _GENDER_RE.search(details_text)
    return m.group(1).capitalize() if m else ""


def _table_to_text(table) -> str:
    """
    Flatten a patient table into a single readable text block.
    Column 0 contains the primary data; column 2 holds cancer target dates.
    Duplicate content in columns 1/2 is deduplicated.
    """
    seen: set[str] = set()
    parts: list[str] = []

    for row in table.rows:
        cells = row.cells
        # Collect unique cell texts (col 0 and col 2 only; col 1 == col 0)
        for idx in (0, 2):
            if idx >= len(cells):
                continue
            text = cells[idx].text.strip()
            if text and text not in seen:
                seen.add(text)
                parts.append(text)

    return "\n\n".join(parts)


def _table_to_cells(table) -> list[CellRef]:
    """Return all cells in the table as a flat list with stable row/col coordinates.

    Duplicate cells are removed using two strategies:
    1. XML element identity (python-docx merged cells share the same _tc)
    2. Adjacent text comparison (Word sometimes copies content without true merge)
    Column indices are renumbered sequentially per row after deduplication.
    """
    cells = []
    for i, row in enumerate(table.rows):
        seen_tcs = set()
        prev_text = None
        col_idx = 0
        for cell in row.cells:
            tc_id = id(cell._tc)
            text = cell.text.strip()
            # Skip if same XML element (true merge) or same text as previous cell
            if tc_id in seen_tcs or (prev_text is not None and text == prev_text):
                continue
            seen_tcs.add(tc_id)
            prev_text = text
            cells.append({"row": i, "col": col_idx, "text": text})
            col_idx += 1
    return cells


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

_MDT_DATE_RE = re.compile(
    r'Multidisciplinary.*?Meeting\s+(\d{2}/\d{2}/\d{4})', re.IGNORECASE
)


_MDT_HEADER_RE = re.compile(
    r'(\w[\w\s]*?)\s*Multidisciplinary.*?Meeting\s+(\d{2}/\d{2}/\d{4})', re.IGNORECASE
)


def _extract_mdt_headers(doc) -> list[dict]:
    """Extract MDT meeting dates and cancer types from paragraph headers.

    The document has one paragraph header per patient like:
    'Colorectal Multidisciplinary Meeting 07/03/2025(i)'
    """
    headers = []
    for para in doc.paragraphs:
        m = _MDT_HEADER_RE.search(para.text)
        if m:
            headers.append({
                "cancer_type": m.group(1).strip(),
                "date": m.group(2)
            })
        elif _MDT_DATE_RE.search(para.text):
            # Fallback: date found but no cancer type prefix
            dm = _MDT_DATE_RE.search(para.text)
            headers.append({
                "cancer_type": "Unknown",
                "date": dm.group(1)
            })
    return headers


def parse_docx(file_path: str) -> list[PatientBlock]:
    """
    Parse a MDT outcome proformas .docx file and return one PatientBlock
    per patient.

    Splitting strategy: each top-level table in the document represents one
    patient (confirmed by document inspection: 50 tables, 50 patients).
    The MDT date comes from the paragraph header before each table.
    """
    doc = Document(file_path)
    patients: list[PatientBlock] = []

    # Extract MDT dates and cancer types from paragraph headers (one per patient)
    mdt_headers = _extract_mdt_headers(doc)

    for idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue  # skip any degenerate table

        details_cell = _clean(rows[1].cells[0].text)
        name = _extract_name(details_cell)
        nhs = _extract_nhs(details_cell)

        # Derive a stable ID: prefer hospital number, fall back to index
        hosp_m = _HOSPITAL_RE.search(details_cell)
        if hosp_m:
            patient_id = _clean(hosp_m.group(1))
        else:
            patient_id = f"PATIENT_{idx + 1:03d}"

        raw_text = _table_to_text(table)
        raw_cells = _table_to_cells(table)

        # Prepend the MDT meeting info to the raw text so the LLM can extract it
        mdt_header = mdt_headers[idx] if idx < len(mdt_headers) else {}
        mdt_date = mdt_header.get("date", "")
        cancer_type = mdt_header.get("cancer_type", "")
        if mdt_date or cancer_type:
            prefix = f"Cancer Type: {cancer_type}\nMDT Meeting Date: {mdt_date}"
            raw_text = f"{prefix}\n\n{raw_text}"
            # Add as synthetic cell so source_cell tracking can find MDT date
            raw_cells.insert(0, {"row": -1, "col": 0, "text": prefix})

        patients.append(PatientBlock(
            id=patient_id,
            initials=_initials(name) if name else "",
            nhs_number=nhs,
            gender=_extract_gender(details_cell),
            mdt_date=mdt_date,
            raw_text=raw_text,
            raw_cells=raw_cells,
        ))

    return patients


def get_raw_text(file_path: str) -> str:
    """
    Debug helper: return the full document text (all paragraphs + all table
    cells) as a single string, useful for exploring document structure.
    """
    doc = Document(file_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        parts.append(_table_to_text(table))

    return "\n\n".join(parts)
